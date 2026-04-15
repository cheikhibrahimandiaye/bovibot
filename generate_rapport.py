"""
Génération du rapport de conception/réalisation BoviBot — 20-30 pages
Auteurs : Papa Ousmane Mané, Cheikh Ibrahima Ndiaye,
          Thierno Abdoulaye Sall, Ndeye Mingue Ndiaye
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Palette ──────────────────────────────────────────────────────────────────
AMBER   = RGBColor(0xC8, 0x7A, 0x10)
DARK    = RGBColor(0x1A, 0x1A, 0x2E)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
GRAY    = RGBColor(0x44, 0x44, 0x55)
GREEN   = RGBColor(0x2E, 0x7D, 0x32)
BLUE    = RGBColor(0x15, 0x65, 0xC0)

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.page_height  = Cm(29.7)
    section.page_width   = Cm(21.0)
    section.left_margin  = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin   = Cm(2.5)
    section.bottom_margin = Cm(2.0)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic
    if color:
        run.font.color.rgb = color

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name  = "Calibri"
        run.font.color.rgb = DARK if level <= 2 else GRAY
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    return p

def para(doc, text="", size=11, bold=False, italic=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p

def code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.8)
    run = p.add_run(code_text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x1A)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:fill'), 'F0F4F0')
    p._p.get_or_add_pPr().append(shading)
    return p

def bullet(doc, text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=10.5)
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = WHITE
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:fill'), '1A1A2E')
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = 'F5F5FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(cell_text)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:fill'), fill)
            cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    doc.add_paragraph()
    return table

def page_break(doc):
    doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 1. PAGE DE GARDE
# ═════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("BoviBot")
run.font.name  = "Calibri"
run.font.size  = Pt(42)
run.font.bold  = True
run.font.color.rgb = DARK

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Gestion d'Élevage Bovin avec IA et PL/SQL")
set_font(run2, size=20, bold=True, color=AMBER)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Rapport de Conception et Réalisation")
set_font(run3, size=15, italic=True, color=GRAY)

doc.add_paragraph()
doc.add_paragraph()

p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_line = p_line.add_run("─" * 40)
run_line.font.color.rgb = AMBER

doc.add_paragraph()

authors = [
    "Papa Ousmane Mané",
    "Cheikh Ibrahima Ndiaye",
    "Thierno Abdoulaye Sall",
    "Ndeye Mingue Ndiaye",
]
for author in authors:
    pa = doc.add_paragraph()
    pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_a = pa.add_run(author)
    set_font(run_a, size=13, bold=True, color=DARK)

doc.add_paragraph()

infos = [
    ("Filière",       "Informatique — Licence 3"),
    ("Établissement", "ESP/UCAD — École Supérieure Polytechnique"),
    ("Cours",         "Intégration IA et Bases de Données Avancées"),
    ("Encadrant",     "Pr. Ahmath Bamba MBACKE"),
    ("Date",          "Avril 2026"),
]
for label, val in infos:
    pi = doc.add_paragraph()
    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = pi.add_run(f"{label} : ")
    set_font(r1, size=11, bold=True, color=GRAY)
    r2 = pi.add_run(val)
    set_font(r2, size=11, color=DARK)

page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES (manuelle)
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, "Table des matières", level=1)
toc = [
    ("1.", "Introduction et contexte métier"),
    ("2.", "Modélisation de la base de données"),
    ("   2.1", "Dictionnaire de données"),
    ("   2.2", "Modèle Conceptuel des Données (MCD)"),
    ("   2.3", "Modèle Logique des Données (MLD)"),
    ("3.", "Éléments PL/SQL"),
    ("   3.1", "Fonctions stockées"),
    ("   3.2", "Procédures stockées"),
    ("   3.3", "Triggers"),
    ("   3.4", "Events MySQL Scheduler"),
    ("4.", "Architecture technique"),
    ("5.", "Prompt Engineering et intégration LLM"),
    ("6.", "Exemples de dialogues LLM"),
    ("7.", "Tests"),
    ("8.", "Guide d'installation et de déploiement"),
    ("9.", "Conclusion et perspectives"),
]
for num, title in toc:
    pt = doc.add_paragraph()
    r1 = pt.add_run(f"{num}  ")
    set_font(r1, size=10.5, bold=True, color=AMBER)
    r2 = pt.add_run(title)
    set_font(r2, size=10.5, color=DARK)
    pt.paragraph_format.space_after = Pt(2)
page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
# 2. INTRODUCTION ET CONTEXTE MÉTIER
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, "1. Introduction et contexte métier", level=1)

heading(doc, "1.1 L'élevage bovin au Sénégal", level=2)
para(doc, """L'élevage constitue l'un des piliers de l'économie rurale sénégalaise. Avec un cheptel bovin estimé à plus de 3,5 millions de têtes, le Sénégal dispose d'un potentiel considérable pour la production de lait, de viande et de travaux agricoles. Les principales races élevées sont la Zébu Gobra (race locale adaptée à la sécheresse), la Ndama (trypanotolérante du sud), et des races améliorées issues de croisements avec des races exotiques comme le Brahman ou le Holstein.""")
para(doc, """Malgré ce potentiel, le secteur fait face à de nombreux défis : manque de traçabilité, suivi sanitaire insuffisant, gestion empirique des gestations, et absence d'outils numériques adaptés aux réalités du terrain. La grande majorité des éleveurs tient ses registres sur papier, rendant toute analyse de performance quasi-impossible.""")

heading(doc, "1.2 Objectifs du projet BoviBot", level=2)
para(doc, """BoviBot est une application web de gestion d'élevage bovin intégrant un assistant LLM (Large Language Model) connecté à une base de données MySQL. Elle a été conçue dans le cadre du cours d'intégration de l'IA et des bases de données avancées (Licence 3, ESP/UCAD) pour démontrer la maîtrise de :""")
bullet(doc, "La conception et la mise en œuvre d'une base de données relationnelle MySQL 8.x avec PL/SQL avancé (procédures, fonctions, triggers, events).")
bullet(doc, "L'intégration d'un LLM (Ollama/Mistral) pour permettre des requêtes en langage naturel (Text-to-SQL) et l'exécution d'actions via une interface conversationnelle.")
bullet(doc, "Le développement d'un backend REST avec FastAPI (Python) et d'un frontend HTML/CSS/JavaScript moderne.")
para(doc, """L'application couvre les modules fonctionnels suivants : suivi du troupeau, pesées et calcul du GMQ, santé vétérinaire, reproduction, alimentation, ventes, et alertes automatisées.""")
page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
# 3. MODÉLISATION
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, "2. Modélisation de la base de données", level=1)

heading(doc, "2.1 Dictionnaire de données", level=2)
para(doc, "Le dictionnaire ci-dessous décrit les tables principales, leurs attributs, types et contraintes.", size=10.5)

# Table: races
para(doc, "Table : races", bold=True, size=10.5)
add_table(doc,
    ["Attribut", "Type", "Contrainte", "Description"],
    [
        ["id", "INT UNSIGNED", "PK, AUTO_INCREMENT", "Identifiant unique"],
        ["nom", "VARCHAR(100)", "NOT NULL, UNIQUE", "Nom de la race (ex: Zébu Gobra)"],
        ["origine", "VARCHAR(100)", "NULL", "Pays/région d'origine"],
        ["poids_adulte_moyen_kg", "DECIMAL(6,2)", "NULL", "Référence de poids adulte en kg"],
        ["created_at", "DATETIME", "NOT NULL, DEFAULT NOW()", "Horodatage de création"],
    ],
    col_widths=[3.5, 3.5, 4.5, 5.0]
)

para(doc, "Table : animaux", bold=True, size=10.5)
add_table(doc,
    ["Attribut", "Type", "Contrainte", "Description"],
    [
        ["id", "INT UNSIGNED", "PK, AUTO_INCREMENT", "Identifiant unique"],
        ["numero_tag", "VARCHAR(20)", "NOT NULL, UNIQUE", "Tag physique (ex: TAG-001)"],
        ["nom", "VARCHAR(100)", "NULL", "Nom de l'animal"],
        ["race_id", "INT UNSIGNED", "FK → races(id)", "Race de l'animal"],
        ["sexe", "ENUM('M','F')", "NOT NULL", "Sexe"],
        ["date_naissance", "DATE", "NOT NULL", "Date de naissance"],
        ["statut", "ENUM('actif','vendu','mort')", "DEFAULT 'actif'", "Statut courant"],
        ["mere_id", "INT UNSIGNED", "FK → animaux(id)", "Référence à la mère"],
        ["pere_id", "INT UNSIGNED", "FK → animaux(id)", "Référence au père"],
        ["poids_actuel_kg", "DECIMAL(6,2)", "NULL", "Mis à jour par sp_enregistrer_pesee"],
    ],
    col_widths=[3.5, 4.0, 4.0, 5.0]
)

para(doc, "Table : pesees", bold=True, size=10.5)
add_table(doc,
    ["Attribut", "Type", "Contrainte", "Description"],
    [
        ["id", "INT UNSIGNED", "PK", "Identifiant"],
        ["animal_id", "INT UNSIGNED", "FK → animaux(id)", "Animal pesé"],
        ["poids_kg", "DECIMAL(6,2)", "NOT NULL", "Poids mesuré en kg"],
        ["date_pesee", "DATE", "NOT NULL", "Date de la pesée"],
        ["agent", "VARCHAR(100)", "NULL", "Opérateur (humain ou 'BoviBot')"],
    ],
    col_widths=[3.0, 3.5, 4.0, 6.0]
)

para(doc, "Table : sante", bold=True, size=10.5)
add_table(doc,
    ["Attribut", "Type", "Contrainte", "Description"],
    [
        ["id", "INT UNSIGNED", "PK", "Identifiant"],
        ["animal_id", "INT UNSIGNED", "FK → animaux(id)", "Animal concerné"],
        ["type", "VARCHAR(100)", "NOT NULL", "Vaccination, soin, contrôle…"],
        ["date_acte", "DATE", "NOT NULL", "Date de l'acte"],
        ["veterinaire", "VARCHAR(100)", "NULL", "Nom du vétérinaire"],
        ["prochain_rdv", "DATE", "NULL", "Déclenche trg_alerte_vaccination si dépassé"],
    ],
    col_widths=[3.0, 3.5, 4.0, 6.0]
)

para(doc, "Table : ventes", bold=True, size=10.5)
add_table(doc,
    ["Attribut", "Type", "Contrainte", "Description"],
    [
        ["id", "INT UNSIGNED", "PK", "Identifiant"],
        ["animal_id", "INT UNSIGNED", "FK → animaux(id)", "Animal vendu"],
        ["acheteur", "VARCHAR(150)", "NOT NULL", "Nom de l'acheteur"],
        ["telephone", "VARCHAR(20)", "NULL", "Contact acheteur"],
        ["prix_fcfa", "DECIMAL(12,2)", "NOT NULL", "Prix en FCFA"],
        ["poids_vente_kg", "DECIMAL(6,2)", "NULL", "Poids au moment de la vente"],
        ["date_vente", "DATE", "NOT NULL", "Date de la transaction"],
    ],
    col_widths=[3.0, 3.5, 4.0, 6.0]
)

para(doc, "Autres tables : reproduction, alimentation, alertes, historique_statut (voir schéma complet en annexe).", italic=True, size=10)

heading(doc, "2.2 Modèle Conceptuel des Données (MCD)", level=2)
para(doc, "Le MCD identifie les entités et leurs relations sémantiques :")
bullet(doc, "ANIMAL est d'une RACE (n:1)")
bullet(doc, "ANIMAL a plusieurs PESEES (1:n)")
bullet(doc, "ANIMAL a plusieurs ACTES SANITAIRES — table sante (1:n)")
bullet(doc, "ANIMAL est impliqué dans REPRODUCTION en tant que mère ou père (n:m auto-jointure)")
bullet(doc, "ANIMAL a plusieurs RATIONS ALIMENTAIRES (1:n)")
bullet(doc, "ANIMAL peut être l'objet d'une VENTE (1:1 par statut)")
bullet(doc, "Le système génère des ALERTES liées ou non à un ANIMAL (n optionnel)")
bullet(doc, "HISTORIQUE_STATUT trace chaque changement de statut d'un ANIMAL (1:n, alimenté par trigger)")

heading(doc, "2.3 Modèle Logique des Données (MLD)", level=2)
code_block(doc, """races(id PK, nom UNIQUE NOT NULL, origine, poids_adulte_moyen_kg)

animaux(id PK, numero_tag UNIQUE NOT NULL, nom, race_id FK→races,
        sexe, date_naissance, statut, mere_id FK→animaux, pere_id FK→animaux,
        poids_actuel_kg)

pesees(id PK, animal_id FK→animaux, poids_kg, date_pesee, agent)

sante(id PK, animal_id FK→animaux, type, description, date_acte,
      veterinaire, prochain_rdv)

reproduction(id PK, mere_id FK→animaux, pere_id FK→animaux,
             date_saillie, date_velage_prevue, date_velage_reel, statut)

alimentation(id PK, animal_id FK→animaux, type_aliment,
             quantite_kg, cout_unitaire_kg, date_ration)

ventes(id PK, animal_id FK→animaux, acheteur, telephone,
       prix_fcfa, poids_vente_kg, date_vente)

alertes(id PK, type, message, niveau, traitee, animal_id FK→animaux, created_at)

historique_statut(id PK, animal_id FK→animaux, ancien_statut,
                  nouveau_statut, date_changement, modifie_par)""")
page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
# 4. ÉLÉMENTS PL/SQL
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, "3. Éléments PL/SQL", level=1)

# ─── Fonctions ───────────────────────────────────────────────────────────────
heading(doc, "3.1 Fonctions stockées", level=2)

heading(doc, "fn_age_en_mois(p_animal_id)", level=3)
para(doc, "Justification métier : L'âge est une donnée centrale dans toute décision élevage (veau/adulte, seuil d'alerte poids, calcul GMQ fiable). Cette fonction encapsule le calcul TIMESTAMPDIFF pour une réutilisation uniforme dans toutes les requêtes.", size=10.5)
add_table(doc,
    ["Paramètre", "Type", "Description"],
    [["p_animal_id", "INT UNSIGNED", "ID de l'animal à interroger"]],
    col_widths=[4, 4, 8]
)
para(doc, "Retourne : INT — âge en mois entiers | -1 si animal introuvable", italic=True, size=10)
code_block(doc, """CREATE FUNCTION fn_age_en_mois(p_animal_id INT UNSIGNED)
RETURNS INT DETERMINISTIC READS SQL DATA
BEGIN
    DECLARE v_date_naissance DATE;
    SELECT date_naissance INTO v_date_naissance
    FROM animaux WHERE id = p_animal_id LIMIT 1;
    IF v_date_naissance IS NULL THEN RETURN -1; END IF;
    RETURN TIMESTAMPDIFF(MONTH, v_date_naissance, CURDATE());
END""")
para(doc, "Utilisations : triggers (trg_alerte_poids_faible), procédures, dashboard KPI âge moyen, event hebdomadaire.", size=10, italic=True)

heading(doc, "fn_gmq(p_animal_id)", level=3)
para(doc, "Justification métier : Le Gain Moyen Quotidien (GMQ) est l'indicateur clé de performance zootechnique. Il mesure la vitesse de croissance d'un animal et permet de détecter les cas sous-performants. La formule (dernier_poids - premier_poids) / nb_jours est appliquée sur l'historique complet des pesées.", size=10.5)
add_table(doc,
    ["Paramètre", "Type", "Description"],
    [["p_animal_id", "INT UNSIGNED", "ID de l'animal"]],
    col_widths=[4, 4, 8]
)
para(doc, "Retourne : DECIMAL(6,3) en kg/jour | 0.000 si < 2 pesées | -1.000 si animal inexistant", italic=True, size=10)
code_block(doc, """CREATE FUNCTION fn_gmq(p_animal_id INT UNSIGNED)
RETURNS DECIMAL(6,3) DETERMINISTIC READS SQL DATA
BEGIN
    DECLARE v_poids_premier DECIMAL(6,2); DECLARE v_date_premiere DATE;
    DECLARE v_poids_dernier DECIMAL(6,2); DECLARE v_date_derniere DATE;
    -- Première pesée
    SELECT poids_kg, date_pesee INTO v_poids_premier, v_date_premiere
    FROM pesees WHERE animal_id = p_animal_id ORDER BY date_pesee ASC LIMIT 1;
    -- Dernière pesée
    SELECT poids_kg, date_pesee INTO v_poids_dernier, v_date_derniere
    FROM pesees WHERE animal_id = p_animal_id ORDER BY date_pesee DESC LIMIT 1;
    IF v_poids_premier IS NULL OR v_date_premiere = v_date_derniere THEN RETURN 0.000; END IF;
    RETURN ROUND((v_poids_dernier - v_poids_premier) / DATEDIFF(v_date_derniere, v_date_premiere), 3);
END""")
para(doc, "Seuils : GMQ ≥ 0.5 kg/j → excellent (vert), 0.3–0.5 → acceptable (orange), < 0.3 → critique (rouge).", size=10, italic=True)

# ─── Procédures ──────────────────────────────────────────────────────────────
heading(doc, "3.2 Procédures stockées", level=2)

heading(doc, "sp_enregistrer_pesee", level=3)
para(doc, "Justification métier : La pesée est l'acte de terrain le plus fréquent. La procédure garantit une transaction atomique : insertion en base, mise à jour du poids courant, calcul du GMQ et génération d'alerte en un seul appel.", size=10.5)
add_table(doc,
    ["Paramètre", "Type", "Description"],
    [
        ["p_animal_id", "INT UNSIGNED", "ID de l'animal"],
        ["p_poids_kg",  "DECIMAL(6,2)", "Poids mesuré en kg"],
        ["p_date_pesee","DATE",          "Date de la pesée"],
        ["p_agent",     "VARCHAR(100)",  "Opérateur (ex: 'BoviBot')"],
    ],
    col_widths=[4, 4, 8]
)
para(doc, "Séquence d'exécution :", bold=True, size=10.5)
bullet(doc, "1. Vérification statut 'actif' de l'animal (SIGNAL si non-actif)")
bullet(doc, "2. INSERT INTO pesees")
bullet(doc, "3. UPDATE animaux SET poids_actuel_kg")
bullet(doc, "4. Calcul GMQ via fn_gmq()")
bullet(doc, "5. Si veau (âge < 6 mois) ET poids < 60 kg → INSERT INTO alertes (niveau='critique')")
bullet(doc, "6. COMMIT + SELECT résultat (gmq_kg_par_jour, statut_alerte)")
para(doc, "Note : Le trigger trg_alerte_poids_faible se déclenche aussi à l'étape 2 comme double sécurité.", italic=True, size=10)
code_block(doc, """-- Appel LLM (après confirmation) :
CALL sp_enregistrer_pesee(1, 325.0, '2026-04-13', 'BoviBot');
-- Résultat : gmq_kg_par_jour=0.520, statut_alerte='OK'""")

heading(doc, "sp_declarer_vente", level=3)
para(doc, "Justification métier : La vente d'un animal modifie son statut de façon irréversible. La procédure s'assure que l'animal est bien actif, enregistre la transaction et déclenche le trigger de journalisation, le tout dans une transaction ACID.", size=10.5)
add_table(doc,
    ["Paramètre", "Type", "Description"],
    [
        ["p_animal_id",  "INT UNSIGNED",  "ID de l'animal"],
        ["p_acheteur",   "VARCHAR(150)",  "Nom de l'acheteur"],
        ["p_telephone",  "VARCHAR(20)",   "Téléphone (nullable)"],
        ["p_prix_fcfa",  "DECIMAL(12,2)", "Prix en FCFA"],
        ["p_poids_kg",   "DECIMAL(6,2)",  "Poids à la vente (nullable)"],
    ],
    col_widths=[4, 4, 8]
)
para(doc, "Séquence :", bold=True, size=10.5)
bullet(doc, "1. SELECT … FOR UPDATE (verrouillage ligne)")
bullet(doc, "2. Vérification statut : SIGNAL si vendu ou mort")
bullet(doc, "3. INSERT INTO ventes")
bullet(doc, "4. UPDATE animaux SET statut = 'vendu' → déclenche trg_historique_statut")
bullet(doc, "5. COMMIT + SELECT vente_id, détails de confirmation")
code_block(doc, """-- Appel LLM (après confirmation) :
CALL sp_declarer_vente(3, 'Oumar Ba', '+221771234567', 280000, 320.0);
-- Résultat : vente_id=2, nouveau_statut='vendu'""")
page_break(doc)

# ─── Triggers ────────────────────────────────────────────────────────────────
heading(doc, "3.3 Triggers", level=2)

heading(doc, "trg_historique_statut", level=3)
para(doc, "Événement : BEFORE UPDATE ON animaux | Justification : Assurer la traçabilité totale de chaque changement de statut (actif → vendu, actif → mort) pour audit et rapport.", size=10.5)
code_block(doc, """CREATE TRIGGER trg_historique_statut
BEFORE UPDATE ON animaux FOR EACH ROW
BEGIN
    IF OLD.statut != NEW.statut THEN
        INSERT INTO historique_statut
            (animal_id, ancien_statut, nouveau_statut, modifie_par)
        VALUES (OLD.id, OLD.statut, NEW.statut, 'system');
    END IF;
END""")
para(doc, "Test : UPDATE animaux SET statut='vendu' WHERE id=1; → Vérifier SELECT * FROM historique_statut;", italic=True, size=10)

heading(doc, "trg_alerte_vaccination", level=3)
para(doc, "Événement : AFTER INSERT ON sante | Justification : Dès qu'un acte vétérinaire est enregistré avec un prochain_rdv dans le passé, le système génère immédiatement une alerte d'avertissement pour que l'éleveur puisse prendre ses dispositions.", size=10.5)
code_block(doc, """CREATE TRIGGER trg_alerte_vaccination
AFTER INSERT ON sante FOR EACH ROW
BEGIN
    IF NEW.prochain_rdv IS NOT NULL AND NEW.prochain_rdv < CURDATE() THEN
        INSERT INTO alertes (type, message, niveau, animal_id)
        VALUES ('vaccination_manquee',
                CONCAT('RDV vétérinaire dépassé pour ', @tag,
                       ' — ', DATEDIFF(CURDATE(), NEW.prochain_rdv), ' jours de retard'),
                'avertissement', NEW.animal_id);
    END IF;
END""")
para(doc, "Test : Insérer sante avec prochain_rdv='2025-01-01' → alerte niveau='avertissement' créée.", italic=True, size=10)

heading(doc, "trg_alerte_poids_faible", level=3)
para(doc, "Événement : AFTER INSERT ON pesees | Justification : Les veaux de moins de 6 mois doivent peser au moins 60 kg. En dessous, ils risquent la malnutrition sévère. Ce trigger génère une alerte critique immédiate pour intervention d'urgence.", size=10.5)
code_block(doc, """CREATE TRIGGER trg_alerte_poids_faible
AFTER INSERT ON pesees FOR EACH ROW
BEGIN
    DECLARE v_age_mois INT;
    SET v_age_mois = fn_age_en_mois(NEW.animal_id);
    IF v_age_mois < 6 AND NEW.poids_kg < 60.0 THEN
        INSERT INTO alertes (type, message, niveau, animal_id)
        VALUES ('poids_critique',
                CONCAT('ALERTE CRITIQUE — veau ', @tag,
                       ' : ', NEW.poids_kg, ' kg à ', v_age_mois, ' mois'),
                'critique', NEW.animal_id);
    END IF;
END""")
para(doc, "Test : CALL sp_enregistrer_pesee(6, 48.0, CURDATE(), 'Test'); → alerte critique générée.", italic=True, size=10)

# ─── Events ──────────────────────────────────────────────────────────────────
heading(doc, "3.4 Events MySQL Scheduler", level=2)
para(doc, "Prérequis : SET GLOBAL event_scheduler = ON; (ou event_scheduler=ON dans my.cnf)", bold=True, size=10.5)

heading(doc, "evt_alerte_velages (quotidien)", level=3)
para(doc, "Justification métier : Les vêlages nécessitent une présence humaine (assistance, soins du veau). L'event envoie une alerte J-7 pour permettre à l'éleveur de se préparer.", size=10.5)
code_block(doc, """CREATE EVENT evt_alerte_velages
ON SCHEDULE EVERY 1 DAY STARTS CURDATE()
DO BEGIN
    -- Supprimer les anciennes alertes velage_imminent non traitées
    DELETE FROM alertes WHERE type = 'velage_imminent' AND traitee = 0;
    -- Créer nouvelles alertes pour vêlages dans les 7 jours
    INSERT INTO alertes (type, message, niveau, animal_id)
    SELECT 'velage_imminent',
           CONCAT('Vêlage prévu dans ', DATEDIFF(r.date_velage_prevue, CURDATE()),
                  ' jours pour mère ID ', r.mere_id),
           'avertissement', r.mere_id
    FROM reproduction r
    WHERE r.statut = 'en_cours'
      AND r.date_velage_prevue BETWEEN CURDATE() AND DATE_ADD(CURDATE(), INTERVAL 7 DAY);
END""")

heading(doc, "evt_rapport_croissance (hebdomadaire)", level=3)
para(doc, "Justification métier : Un résumé hebdomadaire des performances de croissance permet à l'éleveur de suivre l'évolution de son troupeau sans interroger manuellement la base.", size=10.5)
code_block(doc, """CREATE EVENT evt_rapport_croissance
ON SCHEDULE EVERY 1 WEEK STARTS CURDATE()
DO BEGIN
    -- Insérer une ligne de rapport par animal actif
    INSERT INTO alertes (type, message, niveau, animal_id)
    SELECT 'rapport_croissance',
           CONCAT('RAPPORT HEBDO | ', a.numero_tag,
                  ' | Âge: ', fn_age_en_mois(a.id), ' mois',
                  ' | Poids: ', COALESCE(a.poids_actuel_kg, '?'), ' kg',
                  ' | GMQ: ', ROUND(fn_gmq(a.id), 3), ' kg/j'),
           'info', a.id
    FROM animaux a WHERE a.statut = 'actif';
END""")
page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
# 5. ARCHITECTURE TECHNIQUE
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, "4. Architecture technique", level=1)
para(doc, "BoviBot suit une architecture 3-tiers classique enrichie d'une couche LLM :")

heading(doc, "4.1 Vue d'ensemble", level=2)
add_table(doc,
    ["Couche", "Technologie", "Rôle"],
    [
        ["Frontend",    "HTML5 / CSS3 / JavaScript (Vanilla)", "Interface utilisateur : dashboard, chat, CRUD"],
        ["Backend API", "Python 3.12 + FastAPI (port 8002)",   "REST API, orchestration LLM, sécurité"],
        ["LLM",         "Ollama (Mistral 7B / LLaMA3)",        "Text-to-SQL + extraction paramètres actions"],
        ["Base de données", "MySQL 8.x",                       "PL/SQL, triggers, events, pool de connexions"],
    ],
    col_widths=[3.5, 5.5, 7.5]
)

heading(doc, "4.2 Structure du backend", level=2)
code_block(doc, """backend/
├── main.py           ← Point d'entrée FastAPI + pool MySQL
├── config.py         ← Variables d'environnement (.env)
├── database.py       ← Pool MySQL, execute_query/procedure/write, _fix_str UTF-8
├── models.py         ← Modèles Pydantic (ChatRequest, PendingAction…)
├── routers/
│   ├── chat.py       ← POST /chat — orchestration LLM (consultation/action)
│   ├── stats.py      ← GET /stats, /stats/animaux
│   ├── ventes.py     ← GET/POST /ventes
│   ├── pesees.py     ← GET/POST /pesees
│   ├── sante.py      ← GET/POST /sante
│   └── alertes.py    ← GET /alertes, PATCH /alertes/{id}/traiter
└── llm/
    ├── agent.py      ← call_llm(), handle_consultation(), build_pending_action()
    └── prompts.py    ← SYSTEM_PROMPT + DB_SCHEMA""")

heading(doc, "4.3 Flux de données", level=2)
code_block(doc, """CONSULTATION :
  User → POST /chat {message} → call_llm() → SQL généré
       → execute_query(sql) → MySQL → résultats → réponse naturelle

ACTION (pesée/vente) :
  User → POST /chat → LLM extrait paramètres → ACTION_PENDING
       → Frontend affiche modale confirmation
       → User: "Oui" → POST /chat → execute_confirmed_action()
       → execute_procedure("sp_...") → MySQL → résultat

DIRECT (formulaire UI) :
  User → POST /ventes ou /pesees ou /sante
       → resolve_animal_by_tag() → execute_procedure() ou execute_write()
       → MySQL → 200 OK""")

heading(doc, "4.4 Gestion des connexions MySQL", level=2)
para(doc, "Un pool de 5 connexions MySQL (MySQLConnectionPool) est initialisé au démarrage de l'application. Chaque requête emprunte une connexion et la restitue au pool après usage. Le charset utf8mb4 est configuré pour supporter les caractères UTF-8 étendus (accents, caractères locaux).", size=10.5)
para(doc, "Problème rencontré : Le C-extension de mysql-connector-python décodait les bytes UTF-8 stockés en base en latin1, produisant un mojibake (ex : 'ZÃ©bu' au lieu de 'Zébu'). Solution : fonction _fix_str() qui ré-encode la chaîne en latin1 pour retrouver les bytes bruts UTF-8.", size=10, italic=True)
page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
# 6. PROMPT ENGINEERING
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, "5. Prompt Engineering et intégration LLM", level=1)

heading(doc, "5.1 Modèle LLM utilisé", level=2)
para(doc, "BoviBot utilise Ollama (serveur LLM local) avec le modèle Mistral 7B (Instruct). Ce modèle est choisi pour sa capacité à générer du JSON structuré et à suivre des instructions complexes. La température est fixée à 0.1 pour une sortie déterministe.", size=10.5)

heading(doc, "5.2 Structure du System Prompt", level=2)
para(doc, "Le prompt système est composé de 3 blocs :", size=10.5)
bullet(doc, "Bloc 1 — Schéma de base de données : description textuelle de toutes les tables et de leurs colonnes, ainsi que les signatures des fonctions et procédures disponibles.")
bullet(doc, "Bloc 2 — Modes de fonctionnement : règles strictes pour distinguer CONSULTATION (SELECT), ACTION (CALL procédure) et CONVERSATION (réponse générale).")
bullet(doc, "Bloc 3 — Format de réponse JSON obligatoire : le LLM doit répondre UNIQUEMENT en JSON valide, sans texte parasite.")
code_block(doc, """SYSTEM PROMPT (extrait) :
Tu es BoviBot, assistant intelligent de gestion d'élevage bovin.
[SCHÉMA DB]
TABLE animaux : id, numero_tag, nom, race_id, sexe, date_naissance,
                statut (actif|vendu|mort), mere_id, pere_id, poids_actuel_kg
TABLE pesees  : id, animal_id, poids_kg, date_pesee, agent
...
FONCTIONS : fn_age_en_mois(animal_id), fn_gmq(animal_id)
PROCÉDURES: sp_enregistrer_pesee(animal_id, poids_kg, date_pesee, agent)
            sp_declarer_vente(animal_id, acheteur, telephone, prix_fcfa, poids_vente_kg)

[MODES]
MODE CONSULTATION : génère SQL SELECT
MODE ACTION       : extrait paramètres → confirmation → CALL procédure
MODE CONVERSATION : réponse naturelle

[FORMAT JSON OBLIGATOIRE]
Pour CONSULTATION : {"mode":"CONSULTATION","sql":"SELECT...","natural_response":"..."}
Pour ACTION       : {"mode":"ACTION_PENDING","procedure":"sp_...","extracted_params":{...},...}
Pour CONVERSATION : {"mode":"CONVERSATION","natural_response":"..."}""")

heading(doc, "5.3 Gestion du mode CONSULTATION", level=2)
para(doc, "Le LLM génère une requête SELECT en respectant les règles suivantes :", size=10.5)
bullet(doc, "Utilise TOUJOURS a.statut = 'actif' sauf si l'utilisateur demande explicitement les animaux vendus.")
bullet(doc, "Utilise fn_age_en_mois() et fn_gmq() dans les SELECT pertinents.")
bullet(doc, "Convertit les dates relatives ('aujourd'hui', 'ce mois') en CURDATE() ou MONTH().")
bullet(doc, "Ne génère JAMAIS de DELETE, UPDATE, INSERT, DROP, ALTER.")
para(doc, "Côté backend, un filtre de sécurité vérifie la requête générée avant exécution (blacklist : INSERT, UPDATE, DELETE, DROP, ALTER, TRUNCATE, CALL).", size=10.5)

heading(doc, "5.4 Gestion du mode ACTION", level=2)
para(doc, "Règle absolue du CDC : aucun CALL de procédure sans confirmation explicite de l'utilisateur.", bold=True, size=10.5)
para(doc, "Le flux est le suivant :", size=10.5)
bullet(doc, "1. L'utilisateur formule une demande d'action (pesée ou vente).")
bullet(doc, "2. Le LLM retourne mode='ACTION_PENDING' avec les paramètres extraits et un message de confirmation détaillé.")
bullet(doc, "3. Le backend stocke un objet PendingAction en mémoire (clé = session_id).")
bullet(doc, "4. Le frontend affiche une modale de confirmation avec tous les détails.")
bullet(doc, "5. L'utilisateur clique 'Confirmer' → le frontend envoie 'Oui' au chat.")
bullet(doc, "6. Le backend reconnaît 'Oui' comme confirmation, exécute execute_procedure() et retourne ACTION_EXECUTED.")
bullet(doc, "7. Si l'utilisateur répond 'Non' ou 'Annuler', le PendingAction est supprimé sans exécution.")
page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
# 7. EXEMPLES DE DIALOGUES LLM
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, "6. Exemples de dialogues LLM", level=1)

heading(doc, "6.1 Mode CONSULTATION — Requêtes Text-to-SQL", level=2)

para(doc, "Exemple 1 : Liste des animaux actifs avec âge et GMQ", bold=True, size=10.5)
code_block(doc, """Utilisateur : "Liste tous les animaux actifs avec leur âge et GMQ"

SQL généré :
SELECT a.numero_tag, a.nom, r.nom AS race,
       fn_age_en_mois(a.id) AS age_mois,
       ROUND(fn_gmq(a.id), 3) AS gmq_kg_j,
       a.poids_actuel_kg
FROM animaux a
JOIN races r ON r.id = a.race_id
WHERE a.statut = 'actif'
ORDER BY fn_gmq(a.id) DESC;

Réponse BoviBot :
"Voici les 5 animaux actifs. GMQ moyen : 0.482 kg/j.
 TAG-001 (Baaba, Zébu Gobra) — 18 mois — 0.620 kg/j
 TAG-002 (Ndèye, Ndama)       — 24 mois — 0.510 kg/j
 ..." """)

para(doc, "Exemple 2 : Animaux avec GMQ faible", bold=True, size=10.5)
code_block(doc, """Utilisateur : "Quels animaux ont un GMQ inférieur à 0.3 kg/jour ?"

SQL généré :
SELECT a.numero_tag, a.nom, ROUND(fn_gmq(a.id), 3) AS gmq
FROM animaux a
WHERE a.statut = 'actif'
HAVING gmq < 0.3 AND gmq >= 0
ORDER BY gmq ASC;

Réponse BoviBot :
"2 animaux avec GMQ critique : TAG-006 (0.250 kg/j), TAG-007 (0.180 kg/j).
 Ces animaux nécessitent une attention particulière." """)

para(doc, "Exemple 3 : Coût alimentation du mois", bold=True, size=10.5)
code_block(doc, """Utilisateur : "Quel est le coût total d'alimentation ce mois-ci ?"

SQL généré :
SELECT SUM(quantite_kg * cout_unitaire_kg) AS cout_total_fcfa,
       COUNT(DISTINCT animal_id) AS nb_animaux
FROM alimentation
WHERE MONTH(date_ration) = MONTH(CURDATE())
  AND YEAR(date_ration) = YEAR(CURDATE());

Réponse BoviBot :
"Coût total d'alimentation en avril 2026 : 47 500 FCFA pour 5 animaux." """)

para(doc, "Exemple 4 : Femelles vêlant dans les 30 prochains jours", bold=True, size=10.5)
code_block(doc, """Utilisateur : "Quelles femelles vêlent dans les 30 prochains jours ?"

SQL généré :
SELECT a.numero_tag, a.nom,
       r.date_velage_prevue,
       DATEDIFF(r.date_velage_prevue, CURDATE()) AS jours_restants
FROM reproduction r
JOIN animaux a ON a.id = r.mere_id
WHERE r.statut = 'en_cours'
  AND r.date_velage_prevue BETWEEN CURDATE()
      AND DATE_ADD(CURDATE(), INTERVAL 30 DAY)
ORDER BY r.date_velage_prevue ASC;""")

heading(doc, "6.2 Mode ACTION — Enregistrement avec confirmation", level=2)

para(doc, "Exemple 5 : Enregistrement d'une pesée via le chat", bold=True, size=10.5)
code_block(doc, """Utilisateur : "Enregistre une pesée de 325 kg pour TAG-001 aujourd'hui"

LLM identifie (mode ACTION_PENDING) :
  procedure     : sp_enregistrer_pesee
  numero_tag    : TAG-001
  poids_kg      : 325.0
  date_pesee    : 2026-04-13
  agent         : BoviBot

Message de confirmation affiché :
  "Confirmer la pesée ?
     Animal  : TAG-001 / Baaba
     Poids   : 325.0 kg
     Date    : 2026-04-13
     Agent   : BoviBot
   Répondez Oui pour valider."

Utilisateur : "Oui"

Exécution :
  CALL sp_enregistrer_pesee(1, 325.0, '2026-04-13', 'BoviBot')

Réponse BoviBot :
  "Pesée enregistrée. GMQ actuel : 0.620 kg/jour. Statut : OK." """)

para(doc, "Exemple 6 : Déclaration de vente via le chat", bold=True, size=10.5)
code_block(doc, """Utilisateur : "Déclare la vente de TAG-003 à Oumar Ba pour 280 000 FCFA"

LLM identifie (mode ACTION_PENDING) :
  procedure    : sp_declarer_vente
  numero_tag   : TAG-003
  acheteur     : Oumar Ba
  prix_fcfa    : 280000.0
  telephone    : Non renseigné
  poids_vente_kg : Non renseigné

Message de confirmation :
  "Confirmer la vente ?
     Animal   : TAG-003 / Fatou
     Acheteur : Oumar Ba
     Prix     : 280 000 FCFA
     Date     : 2026-04-13 (automatique)
   Répondez Oui pour valider."

Utilisateur : "Oui"

Exécution :
  CALL sp_declarer_vente(3, 'Oumar Ba', NULL, 280000.0, NULL)

Réponse BoviBot :
  "Vente enregistrée (ID #2). Animal TAG-003 désormais marqué 'vendu'." """)
page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
# 8. TESTS
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, "7. Tests", level=1)

heading(doc, "7.1 Cas normaux", level=2)
add_table(doc,
    ["#", "Test", "Action", "Résultat attendu", "Statut"],
    [
        ["T01", "Pesée normale", "CALL sp_enregistrer_pesee(1, 325.0, CURDATE(), 'Test')",
         "Pesée insérée, poids_actuel mis à jour, GMQ calculé", "✅ OK"],
        ["T02", "Vente animal actif", "CALL sp_declarer_vente(2, 'Diop', NULL, 200000, NULL)",
         "Vente insérée, statut→'vendu', historique mis à jour", "✅ OK"],
        ["T03", "Acte sanitaire", "POST /sante {TAG-001, vaccination, 2026-04-13}",
         "Acte inséré, trigger vérifie prochain_rdv", "✅ OK"],
        ["T04", "Consultation GMQ", "Chat: 'Liste animaux actifs avec GMQ'",
         "SQL généré, tableau résultats affiché", "✅ OK"],
        ["T05", "Dashboard KPIs", "GET /stats + /stats/animaux",
         "Animaux actifs, GMQ moyen, vêlages J+7 affichés", "✅ OK"],
        ["T06", "Alertes polling", "GET /alertes?non_traitees_seulement=true",
         "Liste alertes non traitées, badge mis à jour", "✅ OK"],
    ],
    col_widths=[0.8, 3.0, 5.0, 4.5, 1.5]
)

heading(doc, "7.2 Cas limites", level=2)
add_table(doc,
    ["#", "Cas limite", "Action", "Résultat attendu", "Statut"],
    [
        ["TL01", "Vente animal déjà vendu",
         "CALL sp_declarer_vente sur statut='vendu'",
         "SIGNAL SQLSTATE '45000' : 'cet animal a déjà été vendu'", "✅ OK"],
        ["TL02", "Vente animal mort",
         "CALL sp_declarer_vente sur statut='mort'",
         "SIGNAL SQLSTATE '45000' : 'cet animal est décédé'", "✅ OK"],
        ["TL03", "Poids critique veau",
         "CALL sp_enregistrer_pesee(6, 48.0, CURDATE(), 'Test') — veau 2 mois",
         "Alerte critique 'poids_critique' générée automatiquement", "✅ OK"],
        ["TL04", "RDV vétérinaire dépassé",
         "POST /sante avec prochain_rdv='2025-01-01'",
         "trg_alerte_vaccination génère alerte 'avertissement'", "✅ OK"],
        ["TL05", "Animal introuvable",
         "POST /ventes avec numero_tag='TAG-999'",
         "HTTP 404 : 'Animal introuvable : TAG-999'", "✅ OK"],
        ["TL06", "SQL dangereux via chat",
         "Chat: 'DROP TABLE animaux'",
         "Filtre backend bloque : 'requête non autorisée'", "✅ OK"],
        ["TL07", "Pesée — animal non actif",
         "CALL sp_enregistrer_pesee sur animal vendu",
         "SIGNAL SQLSTATE '45000' : statut non actif", "✅ OK"],
        ["TL08", "Confirmation annulée",
         "Chat action → modale → 'Non'",
         "PendingAction supprimé, aucune procédure exécutée", "✅ OK"],
    ],
    col_widths=[0.8, 3.2, 4.5, 4.0, 1.3]
)

heading(doc, "7.3 Tests des Events", level=2)
add_table(doc,
    ["Event", "Test", "Méthode de vérification"],
    [
        ["evt_alerte_velages",
         "Gestation avec date_velage_prevue = CURDATE()+3",
         "Forcer l'event : CALL mysql.event_scheduler.TRIGGER(); puis SELECT * FROM alertes WHERE type='velage_imminent'"],
        ["evt_rapport_croissance",
         "Troupeau de 6 animaux actifs",
         "Vérifier 6 nouvelles alertes type='rapport_croissance' après déclenchement hebdo"],
    ],
    col_widths=[4.0, 5.5, 7.0]
)
page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
# 9. GUIDE D'INSTALLATION
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, "8. Guide d'installation et de déploiement", level=1)

heading(doc, "8.1 Prérequis", level=2)
add_table(doc,
    ["Composant", "Version", "Notes"],
    [
        ["Python",            "3.11+",    "python.org"],
        ["MySQL",             "8.0+",     "mysql.com (ou XAMPP/WAMP)"],
        ["Ollama",            "dernière", "ollama.com — serveur LLM local"],
        ["Mistral 7B",        "latest",   "ollama pull mistral"],
        ["Navigateur",        "moderne",  "Chrome, Firefox, Edge"],
    ],
    col_widths=[4.0, 3.0, 9.5]
)

heading(doc, "8.2 Installation pas à pas", level=2)
para(doc, "Étape 1 — Cloner le dépôt et installer les dépendances Python", bold=True, size=10.5)
code_block(doc, """git clone https://github.com/<votre-compte>/bovibot.git
cd bovibot
pip install -r requirements.txt""")

para(doc, "Étape 2 — Créer la base de données MySQL", bold=True, size=10.5)
code_block(doc, """mysql -u root -p < schema.sql
mysql -u root -p bovibot < triggers_events.sql

-- Vérifications :
SHOW TABLES;  -- 9 tables
SHOW PROCEDURE STATUS WHERE Db='bovibot';  -- 2 procédures
SHOW FUNCTION STATUS WHERE Db='bovibot';   -- 2 fonctions
SHOW TRIGGERS FROM bovibot;               -- 3 triggers
SHOW EVENTS FROM bovibot;                 -- 2 events""")

para(doc, "Étape 3 — Configurer les variables d'environnement", bold=True, size=10.5)
code_block(doc, """# Fichier .env à créer à la racine du projet
DB_HOST=localhost
DB_PORT=3306
DB_USER=bovibot_user
DB_PASSWORD=votre_mot_de_passe
DB_NAME=bovibot
OLLAMA_HOST=http://localhost:11434
OLLAMA_MODEL=mistral
API_PORT=8002""")

para(doc, "Étape 4 — Démarrer Ollama", bold=True, size=10.5)
code_block(doc, """# Télécharger le modèle Mistral (1ère fois seulement)
ollama pull mistral

# Démarrer le serveur Ollama
ollama serve""")

para(doc, "Étape 5 — Lancer le backend FastAPI", bold=True, size=10.5)
code_block(doc, """python -m uvicorn backend.main:app --port 8002 --reload

# Vérification : http://localhost:8002/
# Réponse attendue : {"status":"ok","app":"BoviBot","version":"1.0.0"}""")

para(doc, "Étape 6 — Ouvrir le frontend", bold=True, size=10.5)
code_block(doc, """# Ouvrir directement dans le navigateur :
frontend/index.html

# Ou via un serveur HTTP simple (recommandé) :
cd frontend && python -m http.server 3000
# Puis ouvrir http://localhost:3000""")

heading(doc, "8.3 Déploiement en production", level=2)
add_table(doc,
    ["Composant", "Plateforme recommandée", "Notes"],
    [
        ["Base de données", "PlanetScale / Railway MySQL", "Importer schema.sql + triggers_events.sql"],
        ["Backend FastAPI",  "Railway / Render",           "Dockerfile ou Procfile : uvicorn backend.main:app"],
        ["LLM",             "Serveur VPS + Ollama",        "Ou remplacer par l'API OpenAI (gpt-4o-mini)"],
        ["Frontend",        "Netlify / Vercel / GitHub Pages", "Modifier API_URL vers backend déployé"],
    ],
    col_widths=[3.5, 4.5, 8.5]
)
code_block(doc, """# my.cnf (MySQL persistant) — activer l'event scheduler
[mysqld]
event_scheduler = ON
character-set-server = utf8mb4
collation-server = utf8mb4_unicode_ci""")
page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
# 10. CONCLUSION ET PERSPECTIVES
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, "9. Conclusion et perspectives", level=1)

heading(doc, "9.1 Bilan du projet", level=2)
para(doc, """Le projet BoviBot répond à l'ensemble des exigences du cahier des charges. Les 19 points de conformité ont été validés : les 8 tables de la base de données sont implémentées avec leurs contraintes d'intégrité, les 2 fonctions PL/SQL (fn_age_en_mois, fn_gmq) sont opérationnelles et utilisées dans 5+ contextes différents, les 2 procédures stockées (sp_enregistrer_pesee, sp_declarer_vente) garantissent des transactions ACID avec gestion d'erreurs, les 3 triggers assurent l'automatisation de la journalisation et des alertes, et les 2 events MySQL Scheduler génèrent des rapports périodiques sans intervention humaine.""")
para(doc, """L'intégration du LLM via Ollama/Mistral permet une interaction en langage naturel conforme aux spécifications : mode CONSULTATION (Text-to-SQL sécurisé) et mode ACTION (confirmation obligatoire avant tout CALL de procédure). Le frontend en HTML/CSS/JS pur offre un tableau de bord complet avec 5 pages fonctionnelles, des alertes en temps réel, et une interface de chat IA.""")

heading(doc, "9.2 Difficultés rencontrées", level=2)
bullet(doc, "Encodage UTF-8 (mojibake) : Le C-extension de mysql-connector-python décodait incorrectement les caractères accentués. Solution : fonction _fix_str() avec ré-encodage latin1→utf8.")
bullet(doc, "Gestion des modes LLM : Le modèle Mistral retournait parfois CONVERSATION au lieu d'ACTION_PENDING pour des demandes d'actions. Solution : création d'endpoints REST directs (/ventes, /pesees) contournant le LLM pour les actions formulaires.")
bullet(doc, "Conflits de port (WinError 10013) : Processus uvicorn orphelins bloquant le port 8002. Solution : taskkill /PID.")
bullet(doc, "Nombre d'arguments procédures : sp_declarer_vente installé en DB avec 5 params au lieu de 6 selon schema.sql. Solution : diagnostic via logs, adaptation du code Python.")

heading(doc, "9.3 Perspectives d'évolution", level=2)
bullet(doc, "Export PDF : Génération de rapports PDF des pesées, ventes et bilans sanitaires (feature bonus CDC).")
bullet(doc, "Application mobile : Interface React Native ou PWA pour une utilisation sur le terrain.")
bullet(doc, "Notifications push : Alertes critiques envoyées par SMS (Twilio) ou WhatsApp pour les éleveurs sans accès web permanent.")
bullet(doc, "IA prédictive : Modèle de prédiction du poids adulte basé sur l'historique des pesées et la race (régression linéaire ou Random Forest).")
bullet(doc, "Généalogie interactive : Arbre généalogique visuel des animaux avec leur historique de santé.")
bullet(doc, "Déploiement cloud : Migration vers Railway/Render avec CI/CD GitHub Actions et backup automatique de la base.")
bullet(doc, "Migration vers Claude API : Remplacement d'Ollama local par l'API Anthropic Claude (claude-haiku-4-5) pour de meilleures performances Text-to-SQL en production.")

para(doc, """BoviBot démontre la synergie entre l'intelligence artificielle et les bases de données avancées pour résoudre un problème métier concret. Le projet constitue une base solide pour un déploiement réel dans des élevages sénégalais, avec un potentiel d'impact significatif sur la modernisation du secteur.""", size=11)
page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
# ANNEXE : FICHIER requirements.txt
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, "Annexe — Dépendances Python (requirements.txt)", level=1)
code_block(doc, """fastapi>=0.110.0
uvicorn[standard]>=0.29.0
mysql-connector-python>=8.4.0
pydantic>=2.0.0
pydantic-settings>=2.0.0
ollama>=0.1.0
python-dotenv>=1.0.0""")

heading(doc, "Annexe — Variables de configuration (.env)", level=1)
code_block(doc, """DB_HOST=localhost
DB_PORT=3306
DB_USER=root
DB_PASSWORD=
DB_NAME=bovibot
OLLAMA_HOST=http://localhost:11434
OLLAMA_MODEL=mistral
API_PORT=8002""")

# ── Sauvegarde ────────────────────────────────────────────────────────────────
output_path = "rapport_bovibot.docx"
doc.save(output_path)
print(f"[OK] Rapport genere : {output_path}")
print(f"   Taille : {__import__('os').path.getsize(output_path) // 1024} Ko")
